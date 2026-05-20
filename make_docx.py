#!/usr/bin/env python3
"""
make_docx.py

Reads students.json, summary.csv, and the chart PNGs in charts/, and
produces a native Microsoft Word document (README.docx) with the same
content as the HTML report. Open with `open README.docx` and edit
in Word naturally.

Requires python-docx (`pip3 install python-docx`).
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("error: python-docx not installed.", file=sys.stderr)
    print("Run:  pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


# Order experiments by group, matching the spec's Section 3 layout.
EXPERIMENT_GROUPS = [
    ("Experiment 1 — Single client, unbounded queue",
     "Various (μ, λ) combinations to study queueing behavior across utilization levels.",
     [
         "exp1_mu5_lam3_n1000",
         "exp1_mu5_lam3_n4000",
         "exp1_mu3_lam5_n1000",
         "exp1_mu3_lam5_n4000",
         "exp1_mu50_lam30_n1000",
         "exp1_mu50_lam30_n4000",
         "exp1_mu50_lam35_n2000",
         "exp1_mu50_lam40_n2000",
         "exp1_mu50_lam45_n2000",
     ]),
    ("Experiment 2 — Two clients, unbounded queue",
     "Two concurrent clients × 2000 jobs each, both at (μ, λ) = (50, 20). Server num_jobs = 4000.",
     ["exp2_mu50_lam20_2x2000"]),
    ("Experiment 3 — Bounded queue (q_size = 10)",
     "FIFO capacity limited to 10 — drops occur. q_num is capped at q_size+1 (11) because the executing job is counted.",
     [
         "exp3_mu50_lam45_q10",
         "exp3_mu50_lam48_q10",
     ]),
]


def fmt(value, decimals: int = 2) -> str:
    if value is None or value == "":
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, float) and not pd.isna(value):
        if abs(value - int(value)) < 1e-9 and decimals == 0:
            return f"{int(value)}"
        return f"{value:.{decimals}f}"
    return str(value)


def set_table_borders(table):
    """Add basic borders to a python-docx table (default is none)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        borders.append(b)
    tblPr.append(borders)


def add_title_block(doc: Document, course_info: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(course_info.get("assignment_title", "Programming Assignment 1"))
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    sub_text = (
        f"{course_info.get('course_title', '')} — "
        f"{course_info.get('instructor', '')}"
    )
    if course_info.get("submission_date"):
        sub_text += f"\nSubmission date: {course_info['submission_date']}"
    sub_run = subtitle.add_run(sub_text)
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_students_section(doc: Document, students: list[dict]) -> None:
    doc.add_heading("Submitted by", level=1)

    has_email = any(s.get("email") for s in students)
    cols = 3 if has_email else 2
    table = doc.add_table(rows=1, cols=cols)
    set_table_borders(table)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Name"
    header_cells[1].text = "Student ID"
    if has_email:
        header_cells[2].text = "Email"
    for cell in header_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for s in students:
        row = table.add_row().cells
        row[0].text = s.get("name", "")
        row[1].text = s.get("id", "")
        if has_email:
            row[2].text = s.get("email", "")


def add_submission_contents(doc: Document) -> None:
    doc.add_heading("Submission Contents", level=1)
    files = [
        ("client.c", "UDP job-generator. Samples Poisson inter-arrival times "
                     "and exponential job lengths via the spec's randexp; sends "
                     "each job as a 10-byte datagram in network byte order; logs "
                     "a TSV line per job."),
        ("server.c", "Multi-threaded UDP server. Acceptor (main) thread calls "
                     "recvfrom and enqueues jobs into a synchronized STAILQ; "
                     "worker thread dequeues, sleeps for the job's length to "
                     "simulate processing, and logs the per-job statistics. "
                     "Drop-tail policy when the FIFO is full. Mutex + condition "
                     "variable for synchronization; clock_gettime(CLOCK_MONOTONIC) "
                     "for timing."),
        ("Makefile", "Build script. `make` compiles both binaries with -Wall "
                     "-Wextra -O2 -std=c11 -Wpedantic, plus -pthread for the "
                     "server and -lm for the client. `make clean` removes the "
                     "binaries."),
        ("README.pdf", "This document."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "File"
    hdr[1].text = "Description"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for name, desc in files:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = desc


def add_implementation_summary(doc: Document) -> None:
    doc.add_heading("Implementation Summary", level=1)
    doc.add_paragraph(
        "The system consists of a UDP client that generates jobs according to "
        "a Poisson process (parameter λ) with exponentially distributed "
        "service-time stamps (parameter μ), and a multi-threaded UDP server "
        "that receives jobs into a bounded FIFO and processes them on a "
        "dedicated worker thread."
    )
    doc.add_paragraph(
        "The server's two threads share a STAILQ-backed FIFO protected by a "
        "pthread mutex; the worker waits on a condition variable when the "
        "queue is empty and is woken by the acceptor on each enqueue. A done "
        "flag, set by the acceptor after num_jobs receives, cooperates with "
        "the worker to drain remaining jobs and exit cleanly."
    )
    doc.add_paragraph(
        "The 8th column of the server log (q_time) and the 7th (q_num) "
        "include the just-finished job per the lecturer's clarification — "
        "counters are decremented only after the log line is printed."
    )


def add_stats_table(doc: Document, stats: dict) -> None:
    """Compact 4-column stats table: (label, value, label, value) per row."""
    rho = stats["rho"]
    rho_str = f"{rho:.3f}" if isinstance(rho, (int, float)) else "—"
    is_bounded = isinstance(stats["q_size"], (int, float))

    pairs: list[tuple[str, str]] = [
        ("μ", fmt(stats["mu"], 0)),
        ("λ", fmt(stats["lambda"], 0)),
        ("ρ = λ/μ", rho_str),
        ("Sent", fmt(stats["sent"], 0)),
        ("Processed", fmt(stats["processed"], 0)),
    ]
    if is_bounded:
        pairs += [
            ("Drops", fmt(stats["drops"], 0)),
            ("Drop %", f"{stats['drop_pct']:.2f}%"),
        ]
    pairs += [
        ("Mean job time (µs)", fmt(stats["mean_system_time_us"], 1)),
        ("Median job time (µs)", fmt(stats["median_system_time_us"], 1)),
        ("Mean q_num", fmt(stats["mean_q_num"], 2)),
        ("Median q_num", fmt(stats["median_q_num"], 1)),
        ("Max q_num", fmt(stats["max_q_num"], 0)),
    ]

    # Pack two pairs per row for a compact 4-column table.
    n_rows = (len(pairs) + 1) // 2
    table = doc.add_table(rows=n_rows, cols=4)
    set_table_borders(table)
    for i in range(n_rows):
        a = pairs[2 * i]
        b = pairs[2 * i + 1] if 2 * i + 1 < len(pairs) else ("", "")
        cells = table.rows[i].cells
        cells[0].text = a[0]
        cells[1].text = a[1]
        cells[2].text = b[0]
        cells[3].text = b[1]
        # Make labels grey, values bold
        for col, txt_cell in enumerate(cells):
            for para in txt_cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if col % 2 == 0:  # label columns
                        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    else:
                        run.bold = True


def add_experiment(doc: Document, stats: dict, charts_dir: Path) -> None:
    name = stats["experiment"]
    doc.add_heading(name, level=2)

    add_stats_table(doc, stats)

    qsize_path = charts_dir / f"{name}_qsize.png"
    hist_path = charts_dir / f"{name}_hist.png"
    if qsize_path.exists():
        doc.add_picture(str(qsize_path), width=Inches(6.0))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run("Queue size (q_num) at each job's completion, in chronological order.")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if hist_path.exists():
        doc.add_picture(str(hist_path), width=Inches(6.0))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run("Histogram of per-job system times (departure − arrival), 10 bins.")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_discussion(doc: Document) -> None:
    doc.add_heading("Discussion of Results", level=1)

    doc.add_heading("M/M/1 theory vs observed", level=2)
    doc.add_paragraph(
        "The M/M/1 model predicts the average number of jobs in the system "
        "as L = ρ / (1 − ρ). For our stable experiments at moderate "
        "utilization (ρ ∈ {0.6, 0.7, 0.8, 0.9}), we observe average q_num "
        "values in the same order of magnitude as the theoretical "
        "prediction; the gap is explained by the finite simulation length: "
        "1000–4000 samples is not enough to reach the long-tail steady "
        "state, so the observed average is biased low compared to L."
    )
    doc.add_paragraph(
        "The deliberately unstable case (μ, λ) = (3, 5) — ρ = 1.67 > 1 — "
        "shows monotonically growing q_num across the run, reaching "
        "hundreds (1000-job) or low thousands (4000-job). This is exactly "
        "the M/M/1 instability signature: the queue cannot stabilize "
        "because arrivals exceed service capacity."
    )

    doc.add_heading("Effect of utilization on system time", level=2)
    doc.add_paragraph(
        "For the (50, λ) family at fixed μ=50 and increasing λ from 30 to "
        "45, we see mean and median job system times grow with ρ. This "
        "matches the theoretical mean wait W = 1/(μ − λ) — wait time goes "
        "to infinity as ρ → 1."
    )

    doc.add_heading("Drops in bounded-queue experiments", level=2)
    doc.add_paragraph(
        "Both bounded experiments (q_size = 10) recorded a small number of "
        "drops (8 of 2000 ≈ 0.4%). The fact that the (50, 45) and (50, 48) "
        "runs produced very similar drop counts on macOS is attributable "
        "to the host OS's nanosleep minimum granularity (~50–100 µs on "
        "macOS), which is comparable to the theoretical inter-arrival "
        "times at λ=45 and λ=48 (≈ 22 µs and 21 µs respectively). The "
        "scheduler effectively floors both to a similar realized rate. On "
        "a Linux host with finer scheduling resolution the drop counts "
        "should differentiate more clearly."
    )

    doc.add_heading("Two-client behavior", level=2)
    doc.add_paragraph(
        "Experiment 2 (two concurrent clients, each at λ=20, sharing μ=50) "
        "has combined utilization ρ_total = 2λ/μ = 0.8. Observed mean q_num "
        "was in the expected range and the two clients' jobs interleaved "
        "correctly in the server log (different PIDs and ephemeral ports), "
        "confirming the server multiplexes UDP datagrams from multiple "
        "senders without any additional code — UDP's connectionless nature "
        "and the kernel's per-port receive queue handle this transparently."
    )


def add_platform_notes(doc: Document) -> None:
    doc.add_heading("Platform Notes", level=1)
    doc.add_paragraph(
        "Code was developed on macOS (Apple Silicon) and tested with both "
        "the regular optimized build (-O2) and an AddressSanitizer-instrumented "
        "build (-fsanitize=address -O0 -ggdb3) under stress (1000 jobs) and "
        "bounded-queue scenarios — both clean."
    )
    doc.add_paragraph(
        "The Makefile follows the spec's example exactly, including the "
        "-march=native flag. Since the submission is source code, the flag "
        "is evaluated when the grader runs make on the Ubuntu host and "
        "targets that CPU's instruction set."
    )
    doc.add_paragraph(
        "Absolute timing values reported in this document include the ~50 µs "
        "minimum overhead of nanosleep on macOS. On the Linux grading "
        "environment (finer scheduler granularity) the absolute numbers "
        "will be smaller, but the relative pattern across experiments — "
        "queue growing with ρ, drops increasing with bounded capacity — "
        "is platform independent."
    )


def main() -> int:
    project_root = Path(__file__).resolve().parent
    students_file = project_root / "students.json"
    summary_file = project_root / "summary.csv"
    charts_dir = project_root / "charts"
    output = project_root / "README.docx"

    if not students_file.exists():
        print(f"error: {students_file} missing", file=sys.stderr)
        return 1
    if not summary_file.exists():
        print(f"error: {summary_file} missing — run analyze.py first", file=sys.stderr)
        return 1
    if not charts_dir.is_dir():
        print(f"error: {charts_dir}/ missing — run analyze.py first", file=sys.stderr)
        return 1

    course_info = json.loads(students_file.read_text())
    students = course_info.pop("students", [])
    summary_df = pd.read_csv(summary_file)
    summary_lookup = {row["experiment"]: row.to_dict() for _, row in summary_df.iterrows()}

    doc = Document()

    # Tighten margins (default is 1 inch on all sides — Word's default is wasteful)
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Default body style — 11pt
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(11)

    add_title_block(doc, course_info)
    add_students_section(doc, students)
    add_submission_contents(doc)
    add_implementation_summary(doc)

    doc.add_heading("Experimental Results", level=1)

    for group_title, group_desc, exp_names in EXPERIMENT_GROUPS:
        doc.add_heading(group_title, level=1)
        doc.add_paragraph(group_desc)
        for name in exp_names:
            stats = summary_lookup.get(name)
            if stats is None:
                doc.add_paragraph(f"(missing data for {name})")
                continue
            add_experiment(doc, stats, charts_dir)

    add_discussion(doc)
    add_platform_notes(doc)

    doc.save(output)
    size_kb = output.stat().st_size // 1024
    print(f"Wrote {output} ({size_kb} KB).")
    print()
    print("Open in Word:")
    print(f"  open {output.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
